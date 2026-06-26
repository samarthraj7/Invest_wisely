"""Step 6 - Report assembly + export.

Produces a self-contained HTML memo from the structured report, and attempts a
PDF via WeasyPrint (falls back to HTML if the system lacks the native deps).
"""
from __future__ import annotations

from pathlib import Path

from jinja2 import Template

from ..schemas import InvestmentReport

_SEV_COLOR = {"critical": "#b91c1c", "high": "#dc2626", "medium": "#d97706", "low": "#65a30d"}

_TEMPLATE = Template(
    """<!doctype html><html><head><meta charset="utf-8">
<title>{{ r.company_snapshot.name }} — Investment Memo</title>
<style>
 body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;color:#1e293b;max-width:820px;
      margin:0 auto;padding:40px;line-height:1.5}
 h1{font-size:26px;margin-bottom:2px} h2{border-bottom:2px solid #e2e8f0;padding-bottom:6px;
    margin-top:34px;font-size:18px;color:#0f172a}
 .muted{color:#64748b} .pill{display:inline-block;padding:2px 10px;border-radius:999px;
    font-size:12px;font-weight:600}
 .src{font-size:11px;color:#64748b} .conf{font-size:10px;text-transform:uppercase;
    letter-spacing:.04em;color:#94a3b8;margin-left:6px}
 .claim{margin:6px 0;padding:8px 12px;background:#f8fafc;border-left:3px solid #cbd5e1;border-radius:4px}
 .flag{margin:8px 0;padding:10px 14px;border-radius:6px;background:#fef2f2}
 table{width:100%;border-collapse:collapse;margin-top:8px} td,th{text-align:left;padding:6px 8px;
    border-bottom:1px solid #e2e8f0;font-size:14px;vertical-align:top}
 .note{background:#f1f5f9;border-radius:8px;padding:12px 16px;font-size:13px;color:#475569;margin-top:30px}
 .rec{font-size:20px;font-weight:700}
</style></head><body>

<h1>{{ r.company_snapshot.name }}</h1>
<div class="muted">{{ r.company_snapshot.one_liner }}</div>
<div class="muted">{{ r.company_snapshot.sector or '' }} · {{ r.company_snapshot.stage or '' }}
 · {{ r.company_snapshot.location or '' }} · Ask: {{ r.company_snapshot.ask or 'n/a' }}</div>
{% if r.mock_mode %}<p><span class="pill" style="background:#fef9c3;color:#854d0e">MOCK MODE — add API keys for live research</span></p>{% endif %}
{% if r.warnings %}<div style="background:#fffbeb;border:1px solid #fde68a;border-radius:8px;padding:10px 14px;margin-top:10px">
<b style="color:#854d0e;font-size:13px">Data-quality notes</b>
<ul style="margin:6px 0 0;padding-left:18px;font-size:13px;color:#92400e">
{% for w in r.warnings %}<li>{{ w }}</li>{% endfor %}</ul></div>{% endif %}

{% if r.executive_summary %}<div style="border-left:4px solid #6366f1;background:#eef2ff;
  padding:12px 16px;margin-top:18px;border-radius:6px">
<b style="font-size:12px;text-transform:uppercase;letter-spacing:.04em;color:#4f46e5">Consolidated summary</b>
<p style="margin:6px 0 0;font-size:14px">{{ r.executive_summary }}</p></div>{% endif %}

<h2>1 · Company snapshot</h2>
{% for c in r.company_snapshot.deck_claims %}{{ claim(c) }}{% endfor %}

<h2>2 · Team analysis</h2>
{% for m in r.team_analysis %}
 <p><b>{{ m.name }}</b> — {{ m.title or '' }}
   <span class="conf">research: {{ m.research_confidence.value }}</span></p>
 {% set cr = m.credentials %}
 {% if cr and (cr.assessment or cr.papers_count is not none or cr.patents_count is not none or cr.notable_achievements) %}
 <div style="font-size:12px;color:#475569;margin:2px 0 6px">
   {% if cr.years_experience is not none %}<span class="pill" style="background:#eef2ff;color:#4338ca">~{{ cr.years_experience }}y exp</span> {% endif %}
   {% if cr.papers_count is not none %}<span class="pill" style="background:#ecfeff;color:#0e7490">{{ cr.papers_count }} paper(s){% if cr.research_quality %} · {{ cr.research_quality }}{% endif %}</span> {% endif %}
   {% if cr.patents_count is not none %}<span class="pill" style="background:#f5f3ff;color:#6d28d9">{{ cr.patents_count }} patent(s){% if cr.patent_quality %} · {{ cr.patent_quality }}{% endif %}</span> {% endif %}
   {% if cr.assessment %}<div style="margin-top:4px">{{ cr.assessment }}</div>{% endif %}
 </div>
 {% for c in cr.notable_achievements %}{{ claim(c) }}{% endfor %}
 {% endif %}
 {% if m.researched_background %}<p class="muted" style="font-size:12px;margin:4px 0 0"><b>Background &amp; experience</b></p>
 {% for c in m.researched_background %}{{ claim(c) }}{% endfor %}{% endif %}
 {% if m.strengths %}<p class="muted" style="font-size:12px;margin:4px 0 0"><b>Strengths</b></p>
 {% for c in m.strengths %}{{ claim(c) }}{% endfor %}{% endif %}
 {% if m.founder_market_fit %}<p class="muted" style="font-size:12px;margin:4px 0 0"><b>Founder–market fit</b></p>
 {% for c in m.founder_market_fit %}{{ claim(c) }}{% endfor %}{% endif %}
 {% if m.gaps_vs_venture %}<p class="muted" style="font-size:12px;margin:4px 0 0"><b>Gaps vs. venture</b></p>
 {% for c in m.gaps_vs_venture %}{{ claim(c) }}{% endfor %}{% endif %}
{% endfor %}

<h2>3 · Competitive landscape</h2>
{% for c in r.competitive_landscape.named_in_deck %}{{ claim(c.note) }}{% endfor %}
{% for c in r.competitive_landscape.discovered %}{{ claim(c.note) }}{% endfor %}
{% for c in r.competitive_landscape.differentiation_assessment %}{{ claim(c) }}{% endfor %}

<h2>4 · Red flags</h2>
{% for f in r.red_flags %}
 <div class="flag" style="border-left:4px solid {{ sev(f.severity.value) }}">
  <b style="color:{{ sev(f.severity.value) }}">[{{ f.severity.value|upper }}]</b> {{ f.title }}<br>
  <span class="src">{{ f.reasoning.claim }} — {{ f.reasoning.source_ref }}</span></div>
{% endfor %}

<h2>5 · Suggested diligence questions</h2>
<ol>{% for q in r.diligence_questions %}<li>{{ q.question }}
 <span class="src">→ {{ q.targets_gap }}</span></li>{% endfor %}</ol>

<h2>6 · Valuation analysis</h2>
<p><b>Comp-derived range:</b> {{ r.valuation.range_low }} – {{ r.valuation.range_high }}
 &nbsp;|&nbsp; <b>Deck ask:</b> {{ r.valuation.deck_ask }}
 &nbsp;|&nbsp; <b>Multiples:</b> {{ r.valuation.multiples_used }}</p>
<table><tr><th>Comp</th><th>Detail</th><th>Source</th></tr>
 {% for c in r.valuation.comps %}<tr><td>{{ c.company }}</td><td>{{ c.detail }}</td>
   <td class="src">{{ c.source.source_ref }}</td></tr>{% endfor %}</table>
<p class="muted"><b>Assumptions:</b></p>
<ul>{% for a in r.valuation.assumptions %}<li>{{ a }}</li>{% endfor %}</ul>
{% for c in r.valuation.ask_vs_comps %}{{ claim(c) }}{% endfor %}

<h2>7 · Recommendation</h2>
<p class="rec">{{ r.recommendation.recommendation.value|upper }}
 <span class="pill" style="background:#fee2e2;color:#991b1b">Risk: {{ r.recommendation.risk_rating }}</span></p>
<p>{{ r.recommendation.rationale }}</p>
{% if r.recommendation.suggested_check_size %}<p><b>Suggested check:</b>
 {{ r.recommendation.suggested_check_size }}</p>{% endif %}
<p class="muted"><b>Named risk factors:</b></p>
{% for c in r.recommendation.risk_factors %}{{ claim(c) }}{% endfor %}

{% if r.delivery and r.delivery.available %}
<h2>8 · Pitch delivery {% if r.delivery.source %}<span class="conf">({{ r.delivery.source }})</span>{% endif %}</h2>
{% if r.delivery.clarity %}<p><b>Clarity:</b> {{ r.delivery.clarity }}</p>{% endif %}
{% if r.delivery.structure %}<p><b>Structure:</b> {{ r.delivery.structure }}</p>{% endif %}
{% if r.delivery.handling_of_questions %}<p><b>Handling of cross-questions:</b> {{ r.delivery.handling_of_questions }}</p>{% endif %}
{% if r.delivery.tone %}<p><b>Tone (from video):</b> {{ r.delivery.tone }}</p>{% endif %}
{% if r.delivery.strengths %}<p class="muted"><b>Delivery strengths:</b></p>
<ul>{% for x in r.delivery.strengths %}<li>{{ x }}</li>{% endfor %}</ul>{% endif %}
{% if r.delivery.weaknesses %}<p class="muted"><b>Delivery weaknesses:</b></p>
<ul>{% for x in r.delivery.weaknesses %}<li>{{ x }}</li>{% endfor %}</ul>{% endif %}
{% if r.delivery.qa %}<p class="muted"><b>Questions &amp; answers:</b></p>
{% for q in r.delivery.qa %}<div class="claim"><b>Q:</b> {{ q.question }}<br><b>A:</b> {{ q.answer }}
<div class="src">Assessment: {{ q.assessment }}<span class="conf">{{ q.confidence.value }}</span></div></div>{% endfor %}{% endif %}
{% endif %}

<div class="note">{{ r.analyst_note }}</div>
</body></html>""",
)


def _claim_html(c) -> str:
    return (
        f'<div class="claim">{c.claim}'
        f'<div class="src">{c.source_type.value}: {c.source_ref}'
        f'<span class="conf">{c.confidence.value}</span></div></div>'
    )


def render_html(report: InvestmentReport) -> str:
    return _TEMPLATE.render(
        r=report,
        claim=_claim_html,
        sev=lambda s: _SEV_COLOR.get(s, "#64748b"),
    )


def _export_html(report: InvestmentReport, out_path: Path) -> tuple[Path, str]:
    html_path = out_path.with_suffix(".html")
    html_path.write_text(render_html(report), encoding="utf-8")
    return html_path, "html"


def _export_pdf(report: InvestmentReport, out_path: Path) -> tuple[Path, str]:
    """Render a PDF via WeasyPrint; fall back to HTML if native deps are absent."""
    try:
        from weasyprint import HTML

        pdf_path = out_path.with_suffix(".pdf")
        HTML(string=render_html(report)).write_pdf(str(pdf_path))
        return pdf_path, "pdf"
    except Exception:
        return _export_html(report, out_path)


def _export_docx(report: InvestmentReport, out_path: Path) -> tuple[Path, str]:
    """Build a .docx memo (cross-platform, no native deps). Falls back to HTML."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except Exception:
        return _export_html(report, out_path)

    r = report
    s = r.company_snapshot
    doc = Document()
    doc.add_heading(s.name, level=0)
    doc.add_paragraph(s.one_liner)
    meta = " · ".join(x for x in [s.sector, s.stage, s.location] if x)
    if meta:
        doc.add_paragraph(meta)
    if s.ask:
        doc.add_paragraph(f"Ask: {s.ask}")
    if r.mock_mode:
        doc.add_paragraph("MOCK MODE — add API keys for live research.")
    for w in r.warnings:
        doc.add_paragraph(f"Note: {w}")
    if r.executive_summary:
        doc.add_heading("Consolidated summary", level=1)
        doc.add_paragraph(r.executive_summary)

    def claim_p(c) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c.claim)
        sub = doc.add_paragraph()
        run = sub.add_run(f"    {c.source_type.value}: {c.source_ref}  ({c.confidence.value})")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_heading("1 · Company snapshot", level=1)
    for c in s.deck_claims:
        claim_p(c)

    doc.add_heading("2 · Team analysis", level=1)
    for m in r.team_analysis:
        doc.add_heading(f"{m.name} — {m.title or ''}  (research: {m.research_confidence.value})", level=2)
        cr = m.credentials
        if cr:
            bits = []
            if cr.years_experience is not None:
                bits.append(f"~{cr.years_experience}y experience")
            if cr.papers_count is not None:
                bits.append(f"{cr.papers_count} paper(s)" + (f" ({cr.research_quality})" if cr.research_quality else ""))
            if cr.patents_count is not None:
                bits.append(f"{cr.patents_count} patent(s)" + (f" ({cr.patent_quality})" if cr.patent_quality else ""))
            if bits:
                doc.add_paragraph("Credentials: " + "  ·  ".join(bits))
            if cr.assessment:
                doc.add_paragraph(cr.assessment)
            for c in cr.notable_achievements:
                claim_p(c)
        for label, claims in (
            ("Background & experience", m.researched_background),
            ("Strengths", m.strengths),
            ("Founder-market fit", m.founder_market_fit),
            ("Gaps vs. venture", m.gaps_vs_venture),
        ):
            if claims:
                doc.add_paragraph(label).runs[0].bold = True
                for c in claims:
                    claim_p(c)

    doc.add_heading("3 · Competitive landscape", level=1)
    for c in r.competitive_landscape.named_in_deck + r.competitive_landscape.discovered:
        doc.add_paragraph(f"{c.name} ({c.relationship})")
        claim_p(c.note)
    for c in r.competitive_landscape.differentiation_assessment:
        claim_p(c)

    doc.add_heading("4 · Red flags", level=1)
    for f in r.red_flags:
        doc.add_paragraph(f"[{f.severity.value.upper()}] {f.title}")
        claim_p(f.reasoning)

    doc.add_heading("5 · Suggested diligence questions", level=1)
    for q in r.diligence_questions:
        doc.add_paragraph(f"{q.question}  (→ {q.targets_gap})", style="List Number")

    doc.add_heading("6 · Valuation analysis", level=1)
    doc.add_paragraph(
        f"Comp-derived range: {r.valuation.range_low} – {r.valuation.range_high}  |  "
        f"Deck ask: {r.valuation.deck_ask}  |  Multiples: {r.valuation.multiples_used}"
    )
    for c in r.valuation.comps:
        doc.add_paragraph(f"{c.company}: {c.detail}  [{c.source.source_ref}]", style="List Bullet")
    for a in r.valuation.assumptions:
        doc.add_paragraph(f"Assumption: {a}", style="List Bullet")
    for c in r.valuation.ask_vs_comps:
        claim_p(c)

    doc.add_heading("7 · Recommendation", level=1)
    rec = r.recommendation
    doc.add_paragraph(f"{rec.recommendation.value.upper()}  ·  Risk: {rec.risk_rating}")
    if rec.suggested_check_size:
        doc.add_paragraph(f"Suggested check: {rec.suggested_check_size}")
    doc.add_paragraph(rec.rationale)
    for c in rec.risk_factors:
        claim_p(c)

    d = r.delivery
    if d and d.available:
        doc.add_heading(f"8 · Pitch delivery ({d.source})", level=1)
        for label, val in (
            ("Clarity", d.clarity),
            ("Structure", d.structure),
            ("Handling of cross-questions", d.handling_of_questions),
            ("Tone (from video)", d.tone),
        ):
            if val:
                doc.add_paragraph(f"{label}: {val}")
        for label, items in (("Delivery strengths", d.strengths), ("Delivery weaknesses", d.weaknesses)):
            if items:
                doc.add_paragraph(label).runs[0].bold = True
                for x in items:
                    doc.add_paragraph(x, style="List Bullet")
        if d.qa:
            doc.add_paragraph("Questions & answers").runs[0].bold = True
            for q in d.qa:
                doc.add_paragraph(f"Q: {q.question}")
                doc.add_paragraph(f"A: {q.answer}")
                doc.add_paragraph(f"   Assessment: {q.assessment} ({q.confidence.value})")

    doc.add_paragraph(r.analyst_note)

    docx_path = out_path.with_suffix(".docx")
    doc.save(str(docx_path))
    return docx_path, "docx"


def export_report(
    report: InvestmentReport, out_path: str | Path, fmt: str = "pdf"
) -> tuple[Path, str]:
    """Export to 'pdf' | 'docx' | 'html'. Returns (path, actual_format).

    PDF/DOCX gracefully fall back to HTML when their dependencies are missing.
    """
    out_path = Path(out_path)
    fmt = (fmt or "pdf").lower()
    if fmt == "html":
        return _export_html(report, out_path)
    if fmt == "docx":
        return _export_docx(report, out_path)
    return _export_pdf(report, out_path)


# Backwards-compatible alias (used by the CLI).
def export_pdf(report: InvestmentReport, out_path: str | Path) -> tuple[Path, str]:
    return export_report(report, out_path, "pdf")
